"""
The flask application for our program
"""
# importing required python libraries
from flask import Flask, jsonify, request, send_file, redirect, url_for, session
from flask_mongoengine import MongoEngine
from flask_cors import CORS, cross_origin
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
from itertools import islice
from webdriver_manager.chrome import ChromeDriverManager
from bson.json_util import dumps
import pandas as pd
import json
from datetime import datetime, timedelta
import yaml
import hashlib
import PyPDF2
import uuid
import io
import os
import openai 
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import certifi
import requests
from email.message import EmailMessage
#import schedule
from threading import Thread
import time
import smtplib
# from google.oauth2 import service_account
# from googleapiclient.discovery import build
# from authlib.integrations.flask_client import OAuth
# from authlib.common.security import generate_token
from fake_useragent import UserAgent


from dotenv import load_dotenv
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

existing_endpoints = ["/applications", "/resume","/recommend","/openai-interact","/set_reminder"]


def reminder_email():
    try:
        # time.sleep(10)
        server = smtplib.SMTP_SSL("smtp.gmail.com",465)
        sender_email = "burnoutapp2023@gmail.com"
        sender_password = "jgny mtda gguq shnw"
        server.login(sender_email,sender_password)

        mongodata = Users.objects()
        for md in mongodata:
            email = md['email']
            for appl in md['applications']:
                if  appl['reminder'] == 1:
                    # message = 'Subject: Daily Reminder to Exercise'
                    # server.sendmail(sender_email,email,message)
                    
                    body = 'Dear '+str(md['fullName'])+'. This is a reminder to apply to '+str(appl['companyName'])+' as '+str(appl['jobTitle'])+'.\n\nThis is an autogenerated reply.\n\nThanks & regards!\nJobTracker Team'

                    message = EmailMessage()
                    message['Subject'] = 'Application reminder'
                    message['From'] = sender_email
                    message['To'] = email
                    message.set_content(body)
                    
                    # Sending the email
                    server.send_message(message)

        server.quit()        
    except KeyboardInterrupt:
        print("Thread interrupted")


# schedule.every().day.at("08:00").do(reminder_email)

# # Run the scheduler
# def schedule_process():
#     while True:
#         schedule.run_pending()
#         time.sleep(10)

# Thread(target=schedule_process).start()


user_agent = UserAgent()
def create_app():
    """
    Creates a server hosted on localhost

    :return: Flask object
    """
    app = Flask(__name__)
    # make flask support CORS
    CORS(app)


    



    # SCOPES = ['https://www.googleapis.com/auth/calendar']
    # credentials = service_account.Credentials.from_service_account_file(
    #     'app.py',
    #     scopes=SCOPES,
    # )
    # service = build('calendar', 'v3', credentials=credentials)


    @app.errorhandler(404)
    def page_not_found(e):  
        """
        Returns a json object to indicate error 404

        :return: JSON object
        """
        return jsonify({"error": "Not Found"}), 404

    @app.errorhandler(405)
    # pylint: disable=C0103
    def page_not_allowed(e):
        """
        Returns a json object to indicate error 405

        :return: JSON object
        """
        return jsonify({"error": "Method not Allowed"}), 405

    @app.before_request
    def middleware():
        """
        Checks for user authorization tokens and returns message

        :return: JSON object
        """
        try:
            if request.method == "OPTIONS":
                return jsonify({"success": "OPTIONS"}), 200
            if request.path in existing_endpoints:
                headers = request.headers
                try:
                    token = headers["Authorization"].split(" ")[1]
                except:
                    return jsonify({"error": "Unauthorized"}), 401
                userid = token.split(".")[0]
                user = Users.objects(id=userid).first()

                if user is None:
                    return jsonify({"error": "Unauthorized"}), 401

                expiry_flag = False
                for tokens in user["authTokens"]:
                    if tokens["token"] == token:
                        expiry = tokens["expiry"]
                        expiry_time_object = datetime.strptime(
                            expiry, "%m/%d/%Y, %H:%M:%S"
                        )
                        if datetime.now() <= expiry_time_object:
                            expiry_flag = True
                        else:
                            delete_auth_token(tokens, userid)
                        break

                if not expiry_flag:
                    return jsonify({"error": "Unauthorized"}), 401

        except:
            return jsonify({"error": "Internal server error"}), 500
        
    

    def get_token_from_header():
        """
        Evaluates token from the request header

        :return: string
        """
        headers = request.headers
        token = headers["Authorization"].split(" ")[1]
        return token

    def get_userid_from_header():
        """
        Evaluates user id from the request header

        :return: string
        """
        headers = request.headers
        token = headers["Authorization"].split(" ")[1]
        userid = token.split(".")[0]
        return userid

    def delete_auth_token(token_to_delete, user_id):
        """
        Deletes authorization token of the given user from the database

        :param token_to_delete: token to be deleted
        :param user_id: user id of the current active user
        :return: string
        """
        user = Users.objects(id=user_id).first()
        auth_tokens = []
        for token in user["authTokens"]:
            if token != token_to_delete:
                auth_tokens.append(token)
        user.update(authTokens=auth_tokens)

    @app.route("/")
    @cross_origin()
    def health_check():
        return jsonify({"message": "Server up and running"}), 200
    
    @app.route('/users/profile', methods=["GET","POST"])
    def get_profile():
        # Retrieve user's applications from MongoDB using user_id
        # user = users_collection.find_one({'id': user_id})
        # if user:
        print('abcd')
        userid = get_userid_from_header()
        user = Users.objects(id=userid).first()
        applications = user["applications"]
        username = user["username"]
        fullname = user["fullName"]
        email = user["email"]
        skills = user["skills"]
        workex = user["workExperience"]
        edu = user["education"]

        stats = {
            'wishlist': sum(1 for app in applications if app['status'] == '1'),
            'applied': sum(1 for app in applications if app['status'] == '2'),
            'waiting_for_referral': sum(1 for app in applications if app['status'] == '3'),
            'rejected': sum(1 for app in applications if app['status'] == '4'),
            'username': username,
            'fullname': fullname,
            'skills': skills,
            'workExp': workex,
            'edu': edu,
            'email': email,
        }
        return jsonify(stats)
        # else:
        #     return jsonify({'message': 'User not found'}), 404
    

    @app.route("/users/signup", methods=["POST"])
    def sign_up():
        """
        Creates a new user profile and adds the user to the database and returns the message

        :return: JSON object
        """
        try:
            # print(request.data)
            data = json.loads(request.data)
            print(data)
            try:
                _ = data["username"]
                _ = data["password"]
                _ = data["fullName"]
            except:
                return jsonify({"error": "Missing fields in input"}), 400

            username_exists = Users.objects(username=data["username"])
            print("herer")
            if len(username_exists) != 0:
                return jsonify({"error": "Username already exists"}), 400
            password = data["password"]
            password_hash = hashlib.md5(password.encode())
            print("hereawdawdar")
            print(data)
            user = Users(
                id=get_new_user_id(),
                fullName=data["fullName"],
                username=data["username"],
                email=data["email"],
                skills= data["skills"],
                education= data["education"],
                workExperience= data["workExperience"],
                password=password_hash.hexdigest(),
                authTokens=[],
                applications=[],
            )
            print("here12123r")
            user.save()
            print("herer22222")
            return jsonify(user.to_json()), 200
        except Exception as e:
            print(e)
            return jsonify({"error": "Internal server error"}), 500

    

















    
    @app.route("/users/login", methods=["POST"])
    def login():
        """
        Logs in the user and creates a new authorization token and stores in the database

        :return: JSON object with status and message
        """
        try:
            try:
                data = json.loads(request.data)
                _ = data["username"]
                _ = data["password"]
            except:
                return jsonify({"error": "Username or password missing"}), 400
            password_hash = hashlib.md5(data["password"].encode()).hexdigest()
            user = Users.objects(
                username=data["username"], password=password_hash
            ).first()
            if user is None:
                return jsonify({"error": "Wrong username or password"})
            token = str(user["id"]) + "." + str(uuid.uuid4())
            expiry = datetime.now() + timedelta(days=1)
            expiry_str = expiry.strftime("%m/%d/%Y, %H:%M:%S")
            auth_tokens_new = user["authTokens"] + [
                {"token": token, "expiry": expiry_str}
            ]
            user.update(authTokens=auth_tokens_new)
            return jsonify({"token": token, "expiry": expiry_str})
        except:
            return jsonify({"error": "Internal server error"}), 500

    @app.route("/users/logout", methods=["POST"])
    def logout():
        """
        Logs out the user and deletes the existing token from the database

        :return: JSON object with status and message
        """
        try:
            userid = get_userid_from_header()
            user = Users.objects(id=userid).first()
            auth_tokens = []
            incoming_token = get_token_from_header()
            for token in user["authTokens"]:
                if token["token"] != incoming_token:
                    auth_tokens.append(token)
            user.update(authTokens=auth_tokens)

            return jsonify({"success": ""}), 200

        except:
            return jsonify({"error": "Internal server error"}), 500

    # search function
    # params:
    #   -keywords: string
    #________________________________________________________________________________________________________________-
    @app.route("/search")
    def search():
        """
        Searches the web and returns the job postings for the given search filters

        :return: JSON object with job results
        """
        keywords = (
            request.args.get("keywords")
            if request.args.get("keywords")
            else "random_test_keyword"
        )
        salary = request.args.get(
            "salary") if request.args.get("salary") else ""
        keywords = keywords.replace(" ", "+")
        if keywords == "random_test_keyword":
            return json.dumps({"label": str("successful test search")})
        # create a url for a crawler to fetch job information
        if salary:
            url = (
                "https://www.google.com/search?q="
                + keywords
                + "%20salary%20"
                + salary
                + "&ibp=htl;jobs"
            )
        else:
            url = "https://www.google.com/search?q=" + keywords + "&ibp=htl;jobs"

       
        headers = {"User-Agent":
                   #    "Mozilla/5.0 (Macintosh; Intel Mac OS X 13_5_2) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/117.0.0.0 Safari/537.36"
                   user_agent.random,
                   "Referrer": "https://www.google.com/"
                   }

        page = requests.get(url, headers=headers)
        soup = BeautifulSoup(page.text, "html.parser")

        job_data = pd.DataFrame(columns=["jobTitle", "companyName", "location", "date","qualifications", "responsibilities", "benefits"])
        job_divs = soup.find_all("div", class_="PwjeAc")

        for index, div in enumerate(job_divs):
            job_data.at[index, "jobTitle"] = div.find("div", {"class": "BjJfJf PUpOsf"}).text
            job_data.at[index, "companyName"] = div.find("div", {"class": "vNEEBe"}).text
            job_data.at[index, "location"] = div.find("div", {"class": "Qk80Jf"}).text
            job_data.at[index, "date"] = div.find_all("span", {"class": "LL4CDc"}, limit=1)[0].text

            # Collect Job Description Details
            desc_elements = div.find_all("div", {"class": "JxVj3d"})
            for element in desc_elements:
                details = list(desc.text for desc in element.find_all("div", {"class": "nDgy9d"}))
                title = element.find("div", {"class": "iflMsb"}).text
                if details:
                    job_data.at[index, str(title).lower()] = details

        missing_columns = list((job_data.loc[:, job_data.isnull().sum(axis=0).astype(bool)]).columns)

        for column in missing_columns:
            job_data.loc[job_data[column].isnull(), [column]] = job_data.loc[job_data[column].isnull(), column].apply(lambda x: [])
        return jsonify(job_data.to_dict("records"))

#____________________________________________________________________________________________________________________________________________   
    # get data from the CSV file for rendering root page
    @app.route("/applications", methods=["GET"])
    def get_data():
        """
        Gets user's applications data from the database

        :return: JSON object with application data
        """
        try:
            userid = get_userid_from_header()
            user = Users.objects(id=userid).first()
            applications = user["applications"]
            return jsonify(applications)
        except:
            return jsonify({"error": "Internal server error"}), 500

    @app.route("/applications", methods=["POST"])
    def add_application():
        """
        Add a new job application for the user

        :return: JSON object with status and message
        """
        try:
            userid = get_userid_from_header()
            try:
                request_data = json.loads(request.data)["application"]
                _ = request_data["jobTitle"]
                _ = request_data["companyName"]
            except:
                return jsonify({"error": "Missing fields in input"}), 400

            user = Users.objects(id=userid).first()
            current_application = {
                "id": get_new_application_id(userid),
                "jobTitle": request_data["jobTitle"],
                "companyName": request_data["companyName"],
                "date": request_data.get("date"),
                "jobLink": request_data.get("jobLink"),
                "location": request_data.get("location"),
                "status": request_data.get("status", "1"),
                "reminder": 0,
            }
            applications = user["applications"] + [current_application]

            user.update(applications=applications)
            return jsonify(current_application), 200
        except:
            return jsonify({"error": "Internal server error"}), 500

    @app.route("/applications/<int:application_id>", methods=["PUT"])
    def update_application(application_id):
        """
        Updates the existing job application for the user

        :param application_id: Application id to be modified
        :return: JSON object with status and message
        """
        try:
            userid = get_userid_from_header()
            try:
                request_data = json.loads(request.data)["application"]
            except:
                return jsonify({"error": "No fields found in input"}), 400

            user = Users.objects(id=userid).first()
            current_applications = user["applications"]

            if len(current_applications) == 0:
                return jsonify({"error": "No applications found"}), 400
            else:
                updated_applications = []
                app_to_update = None
                application_updated_flag = False
                for application in current_applications:
                    if application["id"] == application_id:
                        app_to_update = application
                        application_updated_flag = True
                        for key, value in request_data.items():
                            application[key] = value
                    updated_applications += [application]
                if not application_updated_flag:
                    return jsonify({"error": "Application not found"}), 400
                user.update(applications=updated_applications)

            return jsonify(app_to_update), 200
        except:
            return jsonify({"error": "Internal server error"}), 500

    @app.route("/applications/<int:application_id>", methods=["DELETE"])
    def delete_application(application_id):
        """
        Deletes the given job application for the user

        :param application_id: Application id to be modified
        :return: JSON object with status and message
        """
        try:
            userid = get_userid_from_header()
            user = Users.objects(id=userid).first()

            current_applications = user["applications"]

            application_deleted_flag = False
            updated_applications = []
            app_to_delete = None
            for application in current_applications:
                if application["id"] != application_id:
                    updated_applications += [application]
                else:
                    app_to_delete = application
                    application_deleted_flag = True

            if not application_deleted_flag:
                return jsonify({"error": "Application not found"}), 400
            user.update(applications=updated_applications)
            return jsonify(app_to_delete), 200
        except:
            return jsonify({"error": "Internal server error"}), 500

    @app.route("/recommend", methods=["GET"])
    def recommend_resume():
        """
        Recommends a list of jobs in fortune 500 companies based on the user's resume using pdf parsing and ChatGPT
        """
        try:
            userid = get_userid_from_header()
            try:
                user = Users.objects(id=userid).first()
                if len(user.resume.read()) == 0:
                    raise FileNotFoundError
                else:
                    user.resume.seek(0)
            except:
                return jsonify({"error": "resume could not be found"}), 400
            
            pdf_content = io.BytesIO(user.resume.read())
            load_pdf = PyPDF2.PdfReader(pdf_content)
            page_content = load_pdf.pages[0].extract_text()
            prompt = "Analyse the resume below and recommend a list of 6 jobs for the user. All the comapanies should be among the fortune 500. The recommendations should be in a json format with company name, job title, and a link to the company career page.Only display the json. Json structure is {jobs: [{job_title:xx,company_name:xx,career_page:xx}]\n\nResume:\n\n" + page_content + "\n\nRecommendation JSON:"
            message = [ {"role": "system", "content": prompt} ]
            chat = openai.ChatCompletion.create( 
            model="gpt-3.5-turbo", messages=message
            ) 
            reply = chat.choices[0].message.content 
            return jsonify(reply), 200
        except:
            return jsonify({"error": "Internal server error"}), 500



    @app.route("/resume", methods=["POST"])
    def upload_resume():
        """
        Uploads resume file or updates an existing resume for the user

        :return: JSON object with status and message
        """
        try:
            userid = get_userid_from_header()
            try:
                file = request.files["file"].read()
            except:
                return jsonify({"error": "No resume file found in the input"}), 400

            user = Users.objects(id=userid).first()
            if not user.resume.read():
                # There is no file
                user.resume.put(file)
                user.save()
                return jsonify({"message": "resume successfully uploaded"}), 200
            else:
                # There is a file, we are replacing it
                user.resume.replace(file)
                user.save()
                return jsonify({"message": "resume successfully replaced"}), 200
        except Exception as e:
            print(e)
            return jsonify({"error": "Internal server error"}), 500

    @app.route("/resume", methods=["GET"])
    def get_resume():
        """
        Retrieves the resume file for the user

        :return: response with file
        """
        try:
            userid = get_userid_from_header()
            try:
                user = Users.objects(id=userid).first()
                if len(user.resume.read()) == 0:
                    raise FileNotFoundError
                else:
                    user.resume.seek(0)
            except:
                return jsonify({"error": "resume could not be found"}), 400

            response = send_file(
                user.resume,
                mimetype="application/pdf",
                attachment_filename="resume.pdf",
                as_attachment=True,
            )
            response.headers["x-filename"] = "resume.pdf"
            response.headers["Access-Control-Expose-Headers"] = "x-filename"
            return response, 200
        except:
            return jsonify({"error": "Internal server error"}), 500
    


    @app.route("/set_reminder/<int:application_id>", methods=["POST"])
    def set_reminder(application_id):
        try:
            userid = get_userid_from_header()
        

            user = Users.objects(id=userid).first()
            current_applications = user["applications"]

            if len(current_applications) == 0:
                return jsonify({"error": "No applications found"}), 400
            
            for application in current_applications:
                if application["id"] == application_id:
                    application["reminder"] = 1
                    user.save()

                    return jsonify({'message': 'Reminder set successfully!'})
        except:

            return jsonify({"error": "Internal server error"}), 200
    
    def generate_pdf(data):
        doc = Document()

        # Set page margins to fit within one page
        sections = doc.sections
        for section in sections:
            section.left_margin = Pt(36)  # 0.5 inch
            section.right_margin = Pt(36)  # 0.5 inch
            section.top_margin = Pt(36)  # 0.5 inch
            section.bottom_margin = Pt(36)  # 0.5 inch

        # Helper function to add heading with format
        def add_heading_with_format(doc, text, font_size=16, is_bold=True):
            p = doc.add_paragraph()
            run = p.add_run(text)
            if is_bold:
                run.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run.font.size = Pt(font_size)

        # Function to add details section
        def add_details_section(doc, section_title, details, is_bold_title=True):
            if section_title:
                add_heading_with_format(doc, section_title, font_size=14, is_bold=True)
            for detail in details:
                for key, value in detail.items():
                    if key == "company":
                        p = doc.add_paragraph()
                        run = p.add_run(value)
                        run.bold = True
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    elif key == "project_title":
                        # Add the value of "project_title" with bold formatting
                        p = doc.add_paragraph()
                        run = p.add_run(value)
                        run.bold = True
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    elif key == "descriptionc":
                        # Add the value of "descriptionc" without "descriptionc" prefix
                        doc.add_paragraph(value, style="List Bullet")
                    elif key != "descriptionc" and key != "level" and key != "extracurricularActivities":
                        if key == "university":
                            # Add the value of "university" with bold formatting and without a bullet
                            p = doc.add_paragraph()
                            run = p.add_run("University: " + value)
                            run.bold = True
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        else:
                            doc.add_paragraph(f"{value}", style="List Bullet")

        # Title
        add_heading_with_format(doc, "Resume", font_size=18, is_bold=True)

        # Contact Information
        add_heading_with_format(doc, "Contact Information", font_size=16, is_bold=True)
        doc.add_paragraph("Name: " + data["name"])
        doc.add_paragraph("Address: " + data["address"])
        doc.add_paragraph("Email: " + data["email"])
        doc.add_paragraph("LinkedIn: " + data["linkedin"])
        doc.add_paragraph("Phone: " + data["phone"])

        # Education section
        add_details_section(doc, "Education", data["education"])

        # Skills section
        skills = data["skills"]
        skills_text = ", ".join(skill["skills"] for skill in skills)
        add_heading_with_format(doc, "Skills", font_size=14, is_bold=True)
        doc.add_paragraph(skills_text, style="List Bullet")

        # Work Experience section
        add_heading_with_format(doc, "Work Experience", font_size=16, is_bold=True)
        for entry in data["workExperience"]:
            add_details_section(doc, "", [entry], is_bold_title=False)  # Removed the "Work Entry" heading

        # Projects section
        add_heading_with_format(doc, "Projects", font_size=16, is_bold=True)
        for project in data["projects"]:
            add_details_section(doc, "", [project], is_bold_title=False)  # Removed repeated "Project" heading

        # Save the document to a .docx file

        word_buffer = BytesIO()
        output_file_path = "generated_resume.docx"
        doc.save(word_buffer)
        word_buffer.seek(0)

        return word_buffer


    @app.route('/resumebuilder', methods=['POST'])
    def form_builder():
        try:
            # Assuming the request data is in JSON format
            data = request.json

            # Log the data (you can customize this part)
            print("Received Form Data:")
            for key, value in data.items():
                print(f"{key}: {value}")

            # Generate PDF
            pdf_data = generate_pdf(data)

            # Send the PDF file as a response
            return send_file(pdf_data, mimetype='application/msword', as_attachment=True,
                            attachment_filename='generated_resume.docx')
        except Exception as e:
            print(f"Error processing form data: {str(e)}")
            return "Error processing form data", 500
    
    @app.route('/openai-interact', methods=['POST'])
    def openai_interact():
        try:
            data = request.get_json()
            prompt = data.get('prompt')
            # Make a request to the OpenAI API using the provided prompt
            #prompt = "Analyse the resume below and recommend a list of 6 jobs for the user. All the comapanies should be among the fortune 500. The recommendations should be in a json format with company name, job title, and a link to the company career page.Only display the json. Json structure is {jobs: [{job_title:xx,company_name:xx,career_page:xx}]\n\nResume:\n\n" + page_content + "\n\nRecommendation JSON:"
            message = [ {"role": "system", "content": prompt} ]
            chat = openai.ChatCompletion.create( 
            model="gpt-3.5-turbo", messages=message
            ) 
            
            
            return jsonify(chat)
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    return app


app = create_app()
with open("application.yml") as f:
    info = yaml.load(f, Loader=yaml.FullLoader)
    username = info["username"]
    password = info["password"]
    app.config["MONGODB_SETTINGS"] = {
        "db": "appTracker",
        "host": "mongodb://localhost:27017/",
    }
db = MongoEngine()
db.init_app(app)


class Users(db.Document):
    """
    Users class. Holds full name, username, password, as well as applications and resumes
    """

    id = db.IntField(primary_key=True)
    fullName = db.StringField()
    username = db.StringField()
    password = db.StringField()
    email = db.StringField()
    skills = db.ListField()
    workExperience = db.ListField()
    education = db.ListField()
    authTokens = db.ListField()
    applications = db.ListField()
    resume = db.FileField()

    def to_json(self):
        """
        Returns the user details in JSON object

        :return: JSON object
        """
        return {"id": self.id, "fullName": self.fullName, "username": self.username}


def get_new_user_id():
    """
    Returns the next value to be used for new user

    :return: key with new user_id
    """
    user_objects = Users.objects()
    if len(user_objects) == 0:
        return 1

    new_id = 0
    for a in user_objects:
        new_id = max(new_id, a["id"])

    return new_id + 1


def get_new_application_id(user_id):
    """
    Returns the next value to be used for new application

    :param: user_id: User id of the active user
    :return: key with new application_id
    """
    user = Users.objects(id=user_id).first()

    if len(user["applications"]) == 0:
        return 1

    new_id = 0
    for a in user["applications"]:
        new_id = max(new_id, a["id"])

    return new_id + 1


reminder_email()


if __name__ == "__main__":
    app.run(debug=True)